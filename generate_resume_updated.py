import time
from pathlib import Path
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from docx import Document

WATCH_FOLDER = r"C:\Users\LENOVO\.n8n-files"
TARGET_FILENAME = "tailored_text.txt"
OUTPUT_FILENAME = "tailored_resume.docx"

class TailoredTextHandler(FileSystemEventHandler):
    def on_created(self, event):
        self._handle(event)

    def on_modified(self, event):
        self._handle(event)

    def _handle(self, event):
        if event.is_directory:
            return
        path = Path(event.src_path)
        if path.name != TARGET_FILENAME:
            return

        # small delay to make sure n8n has finished writing the file
        time.sleep(0.5)

        try:
            text = path.read_text(encoding="utf-8")
        except Exception as e:
            print(f"Failed to read {path}: {e}")
            return

        if not text.strip():
            print("File is empty, skipping.")
            return

        output_path = path.parent / OUTPUT_FILENAME
        convert_to_docx(text, output_path)
        print(f"Converted -> {output_path}")

        # save a numbered copy before removing so the original is preserved
        try:
            idx = 0
            while True:
                archive_path = path.parent / f"tailored_text_{idx}.txt"
                if not archive_path.exists():
                    break
                idx += 1
            path.rename(archive_path)
            print(f"Archived -> {archive_path}")
        except Exception as e:
            print(f"Could not archive source file: {e}")


def is_heading(line: str) -> bool:
    """ALL CAPS line, with or without trailing colon, short-ish (not a full sentence)."""
    stripped = line.rstrip(":")
    if not stripped.isupper():
        return False
    # avoid treating long all-caps sentences as headings
    return len(stripped.split()) <= 6


def is_bullet(line: str) -> bool:
    return line.startswith("o\t") or line.startswith("o ") or line.startswith("•") or line.startswith("-")


def strip_bullet_marker(line: str) -> str:
    for marker in ("o\t", "o ", "•", "- "):
        if line.startswith(marker):
            return line[len(marker):].strip()
    return line.strip()


def convert_to_docx(text: str, output_path: Path):
    doc = Document()
    is_first_line = True

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if is_first_line:
            doc.add_heading(line, level=1)
            is_first_line = False
            continue

        if is_heading(line):
            doc.add_heading(line.rstrip(":"), level=2)
        elif is_bullet(line):
            doc.add_paragraph(strip_bullet_marker(line), style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(output_path)


if __name__ == "__main__":
    event_handler = TailoredTextHandler()
    observer = Observer()
    observer.schedule(event_handler, WATCH_FOLDER, recursive=False)
    observer.start()
    print(f"Watching {WATCH_FOLDER} for {TARGET_FILENAME} ...")

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()