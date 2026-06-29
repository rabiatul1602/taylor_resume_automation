import sys
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_section_heading(doc, title):
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)

def add_content(doc, content):
    for line in content.split('\n'):
        trimmed = line.strip()
        if not trimmed:
            continue
        if trimmed.startswith('- ') or trimmed.startswith('• '):
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(trimmed.lstrip('-•').strip()).font.size = Pt(11)
        else:
            para = doc.add_paragraph()
            para.add_run(trimmed).font.size = Pt(11)
            para.paragraph_format.space_after = Pt(3)

def parse_sections(text):
    sections = []
    current_title = None
    current_lines = []
    for line in text.split('\n'):
        match = re.match(r'^([A-Z][A-Z\s&]+):$', line.strip())
        if match:
            if current_title:
                sections.append({'title': current_title, 'content': '\n'.join(current_lines).strip()})
            current_title = match.group(1).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_title:
        sections.append({'title': current_title, 'content': '\n'.join(current_lines).strip()})
    return sections

# Read the tailored resume text from file passed by n8n
input_path = sys.argv[1]
with open(input_path, 'r', encoding='utf-8') as f:
    tailored_text = f.read()

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

sections = parse_sections(tailored_text)
for section in sections:
    add_section_heading(doc, section['title'])
    add_content(doc, section['content'])

output_path = r'C:\Users\LENOVO\.n8n-files\tailored_resume.docx'
doc.save(output_path)
print(f"Saved to {output_path}")